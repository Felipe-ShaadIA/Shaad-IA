import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor

NEGRO = RGBColor(0x00, 0x00, 0x00)

def run_negro(p, texto, size=12, bold=False, underline=False):
    run = p.add_run(texto)
    run.font.name = 'Liberation Serif'
    run.font.size = Pt(size)
    run.font.color.rgb = NEGRO
    run.bold = bold
    run.underline = underline

def agregar_negritas(p, texto, size=12):
    for i, parte in enumerate(re.split(r'\*\*(.*?)\*\*', texto)):
        run = p.add_run(parte)
        run.font.name = 'Liberation Serif'
        run.font.size = Pt(size)
        run.font.color.rgb = NEGRO
        if i % 2 == 1:
            run.bold = True

def guardar_docx(texto, ruta):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Liberation Serif'
    doc.styles['Normal'].font.size = Pt(12)

    for linea in texto.split('\n'):
        linea = linea.strip()
        if not linea:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue
        if re.match(r'^# ', linea):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = Pt(17)
            run_negro(p, linea[2:].strip(), bold=True, underline=True)
        elif re.match(r'^## ', linea):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(17)
            run_negro(p, linea[3:].strip(), bold=True, underline=True)
        elif re.match(r'^### ', linea):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(17)
            texto_ap = linea[4:].strip()
            m = re.match(r'^([^:]+):(.*)', texto_ap)
            if m:
                run_negro(p, m.group(1).strip() + ': ', bold=True, underline=True)
                agregar_negritas(p, m.group(2).strip())
            else:
                run_negro(p, texto_ap, bold=True, underline=True)
        elif re.match(r'^[-] ', linea):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(17)
            agregar_negritas(p, linea[2:].strip())
        elif re.match(r'^\d+\. ', linea):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = Pt(17)
            agregar_negritas(p, re.sub(r'^\d+\. ', '', linea))
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(17)
            agregar_negritas(p, linea)

    doc.save(ruta)