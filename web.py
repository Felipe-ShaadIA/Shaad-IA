import streamlit as st
import anthropic
import base64
import os
import io
import datetime
import requests
import tempfile
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor
import re

try:
    import fitz
    PDF_DISPONIBLE = True
except ImportError:
    PDF_DISPONIBLE = False

st.set_page_config(
    page_title="Shaad IA · Resúmenes Inteligentes",
    page_icon="✦",
    layout="centered",
    initial_sidebar_state="collapsed",
)

FEEDBACK_URL = "https://script.google.com/macros/s/AKfycbzwMjJAxwTqDmwMToomyDf_JIsg3hhmN5l902gp3yKiWSdTSpGuaH6fv-KRtblM9H6IVg/exec"
MODELO = "claude-haiku-4-5"
MAX_FOTOS = 5

CHIPS = [
    "Conceptos clave",
    "Examen mañana",
    "Muy detallado",
    "Definiciones",
    "Resumen corto",
    "Con ejemplos",
]

def icon(name, color="#a855f7", size=16):
    icons = {
        "globe": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="2" y1="12" x2="22" y2="12"/><path d="M12 2a15.3 15.3 0 0 1 4 10 15.3 15.3 0 0 1-4 10 15.3 15.3 0 0 1-4-10 15.3 15.3 0 0 1 4-10z"/></svg>',
        "bar-chart": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/></svg>',
        "message-square": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/></svg>',
        "upload-cloud": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="16 16 12 12 8 16"/><line x1="12" y1="12" x2="12" y2="21"/><path d="M20.39 18.39A5 5 0 0 0 18 9h-1.26A8 8 0 1 0 3 16.3"/></svg>',
        "table": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="18" height="18" rx="2"/><line x1="3" y1="9" x2="21" y2="9"/><line x1="3" y1="15" x2="21" y2="15"/><line x1="9" y1="3" x2="9" y2="21"/></svg>',
        "folder-open": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M22 19a2 2 0 0 1-2 2H4a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2z"/></svg>',
        "settings": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="3"/><path d="M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 0 1-2.83 2.83l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-4 0v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 0 1-2.83-2.83l.06-.06A1.65 1.65 0 0 0 4.68 15a1.65 1.65 0 0 0-1.51-1H3a2 2 0 0 1 0-4h.09A1.65 1.65 0 0 0 4.6 9a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 0 1 2.83-2.83l.06.06A1.65 1.65 0 0 0 9 4.68a1.65 1.65 0 0 0 1-1.51V3a2 2 0 0 1 4 0v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 0 1 2.83 2.83l-.06.06A1.65 1.65 0 0 0 19.4 9a1.65 1.65 0 0 0 1.51 1H21a2 2 0 0 1 0 4h-.09a1.65 1.65 0 0 0-1.51 1z"/></svg>',
        "help-circle": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><path d="M9.09 9a3 3 0 0 1 5.83 1c0 2-3 3-3 3"/><line x1="12" y1="17" x2="12.01" y2="17"/></svg>',
        "clock": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>',
        "star": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="{color}" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polygon points="12 2 15.09 8.26 22 9.27 17 14.14 18.18 21.02 12 17.77 5.82 21.02 7 14.14 2 9.27 8.91 8.26 12 2"/></svg>',
        "user": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2"/><circle cx="12" cy="7" r="4"/></svg>',
        "lock": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="11" width="18" height="11" rx="2" ry="2"/><path d="M7 11V7a5 5 0 0 1 10 0v4"/></svg>',
        "file-text": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/><polyline points="10 9 9 9 8 9"/></svg>',
        "refresh-cw": f'<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}" viewBox="0 0 24 24" fill="none" stroke="{color}" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="23 4 23 10 17 10"/><polyline points="1 20 1 14 7 14"/><path d="M3.51 9a9 9 0 0 1 14.85-3.36L23 10M1 14l4.64 4.36A9 9 0 0 0 20.49 15"/></svg>',
    }
    return icons.get(name, "")

def il(name, color="#a855f7", size=15):
    return f'<span style="display:inline-flex;align-items:center;vertical-align:middle;margin-right:5px">{icon(name, color, size)}</span>'

def get_secret(key, default=""):
    try:
        return st.secrets.get(key, default)
    except:
        return os.environ.get(key, default)

def get_api_key():
    return get_secret("ANTHROPIC_API_KEY")

def get_app_password():
    return get_secret("APP_PASSWORD", "shaad2026")

def get_tema():
    return st.session_state.get("tema", "oscuro")

def get_idioma_app():
    return st.session_state.get("idioma_app", "es")

TEXTOS = {
    "es": {
        "badge": "✦ Especializada en resúmenes",
        "slogan": "Sube tus apuntes y obtén un resumen estructurado en segundos",
        "tab1": "Resumir",
        "tab2": "Historial",
        "tab3": "Ajustes",
        "tab4": "Ayuda",
        "idioma": "Idioma del resumen",
        "nivel": "Nivel de detalle",
        "tablas": "Incluir tablas y esquemas",
        "uploader_titulo": "Sube tus apuntes",
        "uploader_sub": "Arrastra aquí tus fotos o PDFs · JPG · PNG · PDF · Máx. 5 archivos",
        "btn_transformar": "Generar resumen con IA",
        "descarga_txt": "Descargar .txt",
        "descarga_docx": "Descargar .docx",
        "sin_resumenes": "Aún no hay resúmenes guardados.",
        "ir_a_resumir": "Generar mi primer resumen",
        "ajustes_titulo": "Configuración",
        "tema_label": "Tema visual",
        "idioma_app_label": "Idioma de la app",
        "guardar": "Guardar ajustes",
        "ajustes_guardados": "Ajustes guardados.",
        "nombre_ph": "Tu nombre",
        "opinion_ph": "¿Qué mejorarías? ¿Qué te gusta?",
        "enviar_fb": "Enviar feedback",
        "fb_ok": "¡Gracias por tu feedback!",
        "fb_error": "Error al enviar.",
        "fb_vacio": "Escribe tu opinión antes de enviar.",
        "footer": f"Hecho con <span style='color:#f43f5e'>♥</span> para estudiantes como tú &nbsp;·&nbsp; <strong style='color:#a855f7'>Shaad IA</strong> © {datetime.datetime.now().year}",
        "login_badge": "Versión beta privada",
        "login_desc": "Introduce la contraseña para acceder a la beta.",
        "contrasena": "Contraseña",
        "contrasena_ph": "Contraseña...",
        "entrar": "Entrar",
        "contrasena_error": "Contraseña incorrecta.",
        "aviso_foto": "Sube al menos una foto o PDF.",
        "aviso_max": f"Máximo {MAX_FOTOS} archivos. Se usarán los primeros {MAX_FOTOS}.",
        "preparando": "Preparando…",
        "analizando": "Analizando imágenes…",
        "generando": "Generando resumen…",
        "listo": "Listo",
        "instrucciones_titulo": "Instrucciones",
        "instrucciones_ph": "Escribe tus instrucciones...",
    },
    "en": {
        "badge": "✦ Specialized in summaries",
        "slogan": "Upload your notes and get a structured summary in seconds",
        "tab1": "Summarize",
        "tab2": "History",
        "tab3": "Settings",
        "tab4": "Help",
        "idioma": "Summary language",
        "nivel": "Detail level",
        "tablas": "Include tables and diagrams",
        "uploader_titulo": "Upload your notes",
        "uploader_sub": "Drag your photos or PDFs here · JPG · PNG · PDF · Max. 5 files",
        "btn_transformar": "Generate summary with AI",
        "descarga_txt": "Download .txt",
        "descarga_docx": "Download .docx",
        "sin_resumenes": "No summaries saved yet.",
        "ir_a_resumir": "Generate my first summary",
        "ajustes_titulo": "Settings",
        "tema_label": "Visual theme",
        "idioma_app_label": "App language",
        "guardar": "Save settings",
        "ajustes_guardados": "Settings saved.",
        "nombre_ph": "Your name",
        "opinion_ph": "What would you improve? What do you like?",
        "enviar_fb": "Send feedback",
        "fb_ok": "Thanks for your feedback!",
        "fb_error": "Error sending.",
        "fb_vacio": "Write your opinion before sending.",
        "footer": f"Made with <span style='color:#f43f5e'>♥</span> for students like you &nbsp;·&nbsp; <strong style='color:#a855f7'>Shaad IA</strong> © {datetime.datetime.now().year}",
        "login_badge": "Private beta version",
        "login_desc": "Enter the password to access the beta.",
        "contrasena": "Password",
        "contrasena_ph": "Password...",
        "entrar": "Enter",
        "contrasena_error": "Incorrect password.",
        "aviso_foto": "Upload at least one photo or PDF.",
        "aviso_max": f"Maximum {MAX_FOTOS} files. Using the first {MAX_FOTOS}.",
        "preparando": "Preparing…",
        "analizando": "Analyzing images…",
        "generando": "Generating summary…",
        "listo": "Done",
        "instrucciones_titulo": "Instructions",
        "instrucciones_ph": "Write your instructions...",
    },
    "gl": {
        "badge": "✦ Especializada en resumos",
        "slogan": "Sube os teus apuntes e obtén un resumo estruturado en segundos",
        "tab1": "Resumir",
        "tab2": "Historial",
        "tab3": "Axustes",
        "tab4": "Axuda",
        "idioma": "Idioma do resumo",
        "nivel": "Nivel de detalle",
        "tablas": "Incluír táboas e esquemas",
        "uploader_titulo": "Sube os teus apuntes",
        "uploader_sub": "Arrastra aquí as túas fotos ou PDFs · JPG · PNG · PDF · Máx. 5 arquivos",
        "btn_transformar": "Xerar resumo con IA",
        "descarga_txt": "Descargar .txt",
        "descarga_docx": "Descargar .docx",
        "sin_resumenes": "Aínda non hai resumos gardados.",
        "ir_a_resumir": "Xerar o meu primeiro resumo",
        "ajustes_titulo": "Configuración",
        "tema_label": "Tema visual",
        "idioma_app_label": "Idioma da app",
        "guardar": "Gardar axustes",
        "ajustes_guardados": "Axustes gardados.",
        "nombre_ph": "O teu nome",
        "opinion_ph": "Que mellorarías? Que che gusta?",
        "enviar_fb": "Enviar feedback",
        "fb_ok": "Grazas polo teu feedback!",
        "fb_error": "Erro ao enviar.",
        "fb_vacio": "Escribe a túa opinión antes de enviar.",
        "footer": f"Feito con <span style='color:#f43f5e'>♥</span> para estudantes coma ti &nbsp;·&nbsp; <strong style='color:#a855f7'>Shaad IA</strong> © {datetime.datetime.now().year}",
        "login_badge": "Versión beta privada",
        "login_desc": "Introduce o contrasinal para acceder á beta.",
        "contrasena": "Contrasinal",
        "contrasena_ph": "Contrasinal...",
        "entrar": "Entrar",
        "contrasena_error": "Contrasinal incorrecto.",
        "aviso_foto": "Sube polo menos unha foto ou PDF.",
        "aviso_max": f"Máximo {MAX_FOTOS} arquivos. Usaranse os primeiros {MAX_FOTOS}.",
        "preparando": "Preparando…",
        "analizando": "Analizando imaxes…",
        "generando": "Xerando resumo…",
        "listo": "Listo",
        "instrucciones_titulo": "Instrucións",
        "instrucciones_ph": "Escribe as túas instrucións...",
    }
}

def t(key):
    return TEXTOS.get(get_idioma_app(), TEXTOS["es"]).get(key, TEXTOS["es"].get(key, key))

def get_css():
    if get_tema() == "claro":
        return """
:root {
    --bg-base: #f8f7ff; --bg-surface: #ede9ff; --bg-card: #ffffff;
    --border: rgba(124,58,237,0.15); --border-strong: rgba(124,58,237,0.35);
    --purple-bright: #7c3aed; --purple-mid: #6d28d9; --purple-soft: rgba(124,58,237,0.08);
    --green-bright: #16a34a; --green-soft: rgba(22,163,74,0.08);
    --text-primary: #1a1a2e; --text-secondary: #4a4a6a; --text-muted: #7070a0;
    --radius-card: 16px; --radius-btn: 10px;
    --shadow-card: 0 4px 24px rgba(0,0,0,0.08), 0 0 0 1px var(--border);
    --shadow-glow-p: 0 0 30px rgba(124,58,237,0.10);
    --shadow-glow-g: 0 0 30px rgba(22,163,74,0.08);
    --font-display: 'Syne', sans-serif; --font-body: 'DM Sans', sans-serif;
    --placeholder: rgba(70,70,120,0.6);
    --drop-bg: rgba(237,233,255,0.8);
}
html, body, [data-testid="stAppViewContainer"] { background-color: var(--bg-base) !important; }
[data-testid="stAppViewContainer"] {
    background-image: radial-gradient(ellipse 60% 40% at 50% -10%, rgba(124,58,237,0.08) 0%, transparent 70%);
}"""
    else:
        return """
:root {
    --bg-base: #0a0a0f; --bg-surface: #0f0f1a; --bg-card: #13131f;
    --border: rgba(138,92,246,0.18); --border-strong: rgba(138,92,246,0.40);
    --purple-bright: #a855f7; --purple-mid: #7c3aed; --purple-soft: rgba(168,85,247,0.12);
    --green-bright: #4ade80; --green-soft: rgba(74,222,128,0.10);
    --text-primary: #f0eeff; --text-secondary: #a89fc0; --text-muted: #7070a0;
    --radius-card: 16px; --radius-btn: 10px;
    --shadow-card: 0 4px 32px rgba(0,0,0,0.45), 0 0 0 1px var(--border);
    --shadow-glow-p: 0 0 40px rgba(168,85,247,0.15);
    --shadow-glow-g: 0 0 40px rgba(74,222,128,0.12);
    --font-display: 'Syne', sans-serif; --font-body: 'DM Sans', sans-serif;
    --placeholder: rgba(200,190,230,0.55);
    --drop-bg: rgba(15,15,26,0.8);
}
html, body, [data-testid="stAppViewContainer"] { background-color: var(--bg-base) !important; }
[data-testid="stAppViewContainer"] {
    background-image:
        radial-gradient(ellipse 60% 40% at 50% -10%, rgba(124,58,237,0.20) 0%, transparent 70%),
        radial-gradient(ellipse 40% 30% at 90% 80%, rgba(34,197,94,0.08) 0%, transparent 60%);
}"""

# Generamos el SVG de la nube para inyectarlo como data-uri en el pseudo-elemento
upload_svg_b64 = base64.b64encode(
    icon("upload-cloud", "#a855f7", 40).encode()
).decode()

titulo_uploader = t("uploader_titulo") if "uploader_titulo" in TEXTOS["es"] else "Sube tus apuntes"
sub_uploader = t("uploader_sub") if "uploader_sub" in TEXTOS["es"] else ""

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;1,9..40,300&display=swap');
{get_css()}
html, body, [data-testid="stAppViewContainer"] {{ font-family: var(--font-body) !important; color: var(--text-primary) !important; }}
[data-testid="stHeader"], [data-testid="stToolbar"], footer {{ display: none !important; }}
.block-container {{ max-width: 720px !important; padding: 2rem 1.5rem 4rem !important; }}

.hero-wrap {{ text-align: center; padding: 1rem 0 0.8rem; }}
.hero-badge {{ display: inline-flex; align-items: center; gap: 6px; background: var(--purple-soft); border: 1px solid var(--border-strong); border-radius: 999px; padding: 4px 14px; font-size: 0.70rem; letter-spacing: 0.12em; text-transform: uppercase; color: var(--purple-bright); font-family: var(--font-display); font-weight: 600; margin-bottom: 0.6rem; }}
.hero-title {{ font-family: var(--font-display) !important; font-size: clamp(1.8rem, 5vw, 2.8rem) !important; font-weight: 800 !important; line-height: 1.05 !important; background: linear-gradient(135deg, #c084fc 0%, #a855f7 35%, #4ade80 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; margin: 0 0 0.3rem !important; letter-spacing: -0.02em; }}
.hero-sub {{ font-family: var(--font-body); font-size: 0.88rem; color: var(--text-secondary); line-height: 1.5; }}

.card {{ background: var(--bg-card); border: 1px solid var(--border); border-radius: var(--radius-card); padding: 1.4rem 1.6rem; box-shadow: var(--shadow-card); margin-bottom: 1rem; }}
.card-label {{ display: flex; align-items: center; gap: 8px; font-family: var(--font-display); font-weight: 700; font-size: 0.78rem; letter-spacing: 0.08em; text-transform: uppercase; color: var(--text-secondary); margin-bottom: 0.8rem; }}
.card-label .dot {{ width: 5px; height: 5px; border-radius: 50%; background: var(--purple-bright); box-shadow: 0 0 6px var(--purple-bright); flex-shrink: 0; }}

/* Pestañas con iconos inline */
[data-testid="stTabs"] [data-baseweb="tab-list"] {{
    background: var(--bg-surface) !important;
    border-radius: 12px !important; padding: 4px !important;
    border: 1px solid var(--border) !important; gap: 2px !important;
}}
[data-testid="stTabs"] [data-baseweb="tab"] {{
    background: transparent !important; color: var(--text-muted) !important;
    border-radius: 8px !important; font-family: var(--font-display) !important;
    font-weight: 600 !important; font-size: 0.82rem !important; border: none !important;
    padding: 6px 12px !important; display: flex !important; align-items: center !important; gap: 6px !important;
}}
[data-testid="stTabs"] [aria-selected="true"] {{
    background: var(--purple-soft) !important; color: var(--purple-bright) !important;
    border: 1px solid var(--border-strong) !important;
}}
[data-testid="stTabs"] [aria-selected="true"] p,
[data-testid="stTabs"] [data-baseweb="tab"] p {{
    color: inherit !important; font-size: inherit !important;
    font-family: inherit !important; font-weight: inherit !important;
    margin: 0 !important;
}}

/* Chips */
[data-testid="stColumn"] [data-testid="stButton"] > button {{
    background: transparent !important;
    border: 1.5px solid #a855f7 !important;
    border-radius: 999px !important;
    color: #f0eeff !important;
    font-family: var(--font-body) !important;
    font-size: 0.80rem !important;
    font-weight: 500 !important;
    padding: 5px 8px !important;
    width: 100% !important;
    box-shadow: none !important;
    outline: none !important;
    transition: all 0.18s !important;
    white-space: nowrap !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
}}
[data-testid="stColumn"] [data-testid="stButton"] > button:hover {{
    background: linear-gradient(135deg, #7c3aed, #a855f7, #06b6d4) !important;
    border-color: transparent !important; color: white !important;
    box-shadow: 0 0 12px rgba(168,85,247,0.35) !important;
}}
[data-testid="stColumn"] [data-testid="stButton"] > button:focus {{
    outline: none !important; box-shadow: none !important;
}}

/* Botón principal */
[data-testid="stButton"] > button[kind="primary"] {{
    background: linear-gradient(135deg, #7c3aed, #a855f7, #06b6d4) !important;
    color: #fff !important; border: none !important; border-radius: 14px !important;
    font-family: var(--font-display) !important; font-weight: 800 !important;
    font-size: 1.05rem !important; padding: 1rem 2rem !important; width: 100% !important;
    box-shadow: 0 0 30px rgba(168,85,247,0.45), 0 4px 20px rgba(124,58,237,0.40) !important;
    transition: transform 0.15s, box-shadow 0.15s !important; outline: none !important;
}}
[data-testid="stButton"] > button[kind="primary"]:hover {{
    transform: translateY(-3px) !important;
    box-shadow: 0 0 50px rgba(168,85,247,0.65), 0 8px 28px rgba(124,58,237,0.55) !important;
}}

/* Inputs */
[data-testid="stTextInput"] input {{
    background: var(--bg-surface) !important; border: 1px solid var(--border) !important;
    border-radius: 10px !important; color: var(--text-primary) !important;
    font-family: var(--font-body) !important; padding: 0.75rem 1rem !important;
}}
[data-testid="stTextInput"] input:focus {{ border-color: var(--purple-bright) !important; box-shadow: 0 0 0 3px rgba(168,85,247,0.15) !important; outline: none !important; }}
[data-testid="stTextInput"] input::placeholder {{ color: var(--placeholder) !important; opacity: 1 !important; }}
[data-testid="stTextInput"] label {{ color: var(--text-secondary) !important; font-size: 0.88rem !important; }}
[data-testid="stTextArea"] textarea {{
    background: var(--bg-surface) !important; border: 1px solid var(--border) !important;
    border-radius: 10px !important; color: var(--text-primary) !important;
    font-family: var(--font-body) !important;
}}
[data-testid="stTextArea"] textarea::placeholder {{ color: var(--placeholder) !important; opacity: 1 !important; }}
[data-testid="stTextArea"] label {{ color: var(--text-secondary) !important; }}

/* Uploader — ocultar label nativo, estilizar dropzone */
[data-testid="stFileUploader"] label {{ display: none !important; }}
[data-testid="stFileUploader"] {{ margin-top: 0.5rem; }}
section[data-testid="stFileUploaderDropzone"] {{
    background: var(--drop-bg) !important;
    border: 2px dashed #a855f7 !important;
    border-radius: var(--radius-card) !important;
    min-height: 160px !important;
    padding: 1.5rem 1rem !important;
    transition: border-color 0.2s, background 0.2s !important;
    position: relative !important;
}}
section[data-testid="stFileUploaderDropzone"]:hover {{
    background: rgba(168,85,247,0.07) !important;
    border-color: #c084fc !important;
}}
/* Centrar todo el contenido de la dropzone */
section[data-testid="stFileUploaderDropzone"] > div {{
    display: flex !important;
    flex-direction: column !important;
    align-items: center !important;
    justify-content: center !important;
    gap: 0.3rem !important;
    width: 100% !important;
    text-align: center !important;
}}
/* Icono nativo de Streamlit — hacerlo más grande y morado */
[data-testid="stFileUploader"] svg {{
    width: 36px !important; height: 36px !important;
    stroke: #a855f7 !important; fill: none !important;
}}
/* Texto principal dentro del uploader */
[data-testid="stFileUploaderDropzoneInstructions"] > div > span {{
    font-family: var(--font-display) !important;
    font-size: 1rem !important;
    font-weight: 700 !important;
    color: var(--text-primary) !important;
    display: block !important;
}}
/* Ocultar el "200MB per file" en inglés */
[data-testid="stFileUploaderDropzoneInstructions"] small {{
    display: none !important;
}}
/* Subtexto de formatos */
[data-testid="stFileUploaderDropzoneInstructions"] > div > small {{
    display: none !important;
}}
/* Botón Browse centrado */
[data-testid="stFileUploaderDropzone"] button {{
    margin: 0.4rem auto 0 !important;
    display: block !important;
    background: var(--purple-soft) !important;
    border: 1px solid var(--border-strong) !important;
    border-radius: 8px !important;
    color: var(--purple-bright) !important;
    padding: 5px 20px !important;
    font-size: 0.82rem !important;
    font-weight: 600 !important;
    cursor: pointer !important;
    transition: all 0.18s !important;
}}
[data-testid="stFileUploaderDropzone"] button:hover {{
    background: var(--purple-bright) !important;
    color: white !important;
}}
[data-testid="stFileUploader"] * {{ color: var(--text-secondary) !important; }}

/* Selectbox */
[data-testid="stSelectbox"] > div > div {{ background: var(--bg-surface) !important; border: 1px solid var(--border) !important; border-radius: 10px !important; color: var(--text-primary) !important; }}
[data-testid="stSelectbox"] label {{ color: var(--text-secondary) !important; font-size: 0.85rem !important; }}
[data-baseweb="select"] * {{ color: var(--text-primary) !important; }}
[data-baseweb="popover"] {{ background: var(--bg-card) !important; border: 1px solid var(--border-strong) !important; border-radius: 12px !important; }}
[role="option"]:hover {{ background: var(--purple-soft) !important; }}

/* Download */
[data-testid="stDownloadButton"] > button {{
    background: var(--bg-card) !important; color: var(--text-primary) !important;
    border: 1px solid var(--border-strong) !important; border-radius: 10px !important;
    font-family: var(--font-display) !important; font-weight: 600 !important;
    font-size: 0.85rem !important; padding: 0.6rem 1.2rem !important; width: 100% !important;
    box-shadow: none !important; transition: all 0.2s !important; outline: none !important;
}}
[data-testid="stDownloadButton"] > button:hover {{
    background: var(--purple-soft) !important; border-color: var(--purple-bright) !important;
    transform: translateY(-1px) !important;
}}

[data-testid="stAlert"] {{ background: var(--bg-card) !important; border-radius: 12px !important; border: 1px solid var(--border) !important; }}
[data-testid="stExpander"] {{ background: var(--bg-card) !important; border: 1px solid var(--border) !important; border-radius: var(--radius-card) !important; }}
[data-testid="stExpander"] > div > div {{ background: var(--bg-card) !important; }}
[data-testid="stExpander"] summary {{ color: var(--text-secondary) !important; font-family: var(--font-display) !important; }}
[data-testid="stMarkdownContainer"] p, [data-testid="stMarkdownContainer"] li {{ color: var(--text-primary) !important; font-family: var(--font-body) !important; line-height: 1.7 !important; }}
[data-testid="stMarkdownContainer"] h1, [data-testid="stMarkdownContainer"] h2, [data-testid="stMarkdownContainer"] h3 {{ font-family: var(--font-display) !important; color: var(--text-primary) !important; }}
[data-testid="stImage"] img {{ border-radius: 12px !important; border: 1px solid var(--border) !important; }}
hr {{ border-color: var(--border) !important; }}
::-webkit-scrollbar {{ width: 6px; }}
::-webkit-scrollbar-track {{ background: var(--bg-base); }}
::-webkit-scrollbar-thumb {{ background: var(--border-strong); border-radius: 3px; }}

.result-card {{ background: linear-gradient(135deg, rgba(74,222,128,0.05) 0%, var(--bg-card) 50%); border: 1px solid rgba(74,222,128,0.25); border-radius: var(--radius-card); padding: 1.6rem 1.8rem; box-shadow: var(--shadow-card), var(--shadow-glow-g); margin-top: 1.2rem; }}
.result-badge {{ background: var(--green-soft); border: 1px solid rgba(74,222,128,0.30); color: var(--green-bright); border-radius: 999px; padding: 3px 12px; font-size: 0.70rem; letter-spacing: 0.08em; text-transform: uppercase; font-family: var(--font-display); font-weight: 600; margin-bottom: 1rem; display: inline-block; }}
.feedback-item {{ background: var(--bg-surface); border: 1px solid var(--border); border-radius: 10px; padding: 0.8rem 1rem; margin-bottom: 0.5rem; }}
.empty-state {{ text-align: center; padding: 3rem 1rem; }}
.footer-custom {{ text-align: center; padding: 2rem 0 0; color: var(--text-muted); font-size: 0.78rem; font-family: var(--font-body); }}

[data-testid="stToggle"] label {{ color: var(--text-secondary) !important; }}
[data-testid="stSlider"] > div > div > div {{ background: linear-gradient(135deg, #7c3aed, #a855f7, #06b6d4) !important; }}
[data-testid="stSlider"] [role="slider"] {{ background: #a855f7 !important; border: 2px solid #c084fc !important; box-shadow: 0 0 12px rgba(168,85,247,0.6) !important; }}
[data-testid="stSlider"] label {{ color: var(--text-secondary) !important; }}
</style>
""", unsafe_allow_html=True)


def encode_image(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=90)
    return base64.b64encode(buf.getvalue()).decode("utf-8")

def pdf_a_imagenes(archivo_pdf):
    imagenes = []
    if not PDF_DISPONIBLE:
        st.error("PDF no disponible.")
        return imagenes
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(archivo_pdf.read())
        tmp_path = tmp.name
    doc = fitz.open(tmp_path)
    for page in doc:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        imagenes.append(img)
    doc.close()
    os.unlink(tmp_path)
    return imagenes

def transcribir(imagenes, idioma, nivel, contexto, esquemas):
    api_key = get_api_key()
    if not api_key:
        st.error("No se encontró la API key.")
        return ""
    client = anthropic.Anthropic(api_key=api_key)
    lineas = {"Breve (2-3 líneas)": "2-3", "Medio (4-6 líneas)": "4-6", "Detallado (7-10 líneas)": "7-10"}.get(nivel, "4-6")
    ctx = f"\nInstrucciones: {contexto}" if contexto.strip() else ""
    sin_esquemas = "\nNO uses tablas, esquemas visuales ni diagramas. Solo texto estructurado." if not esquemas else ""

    system = f"""Eres un experto en crear apuntes de estudio para estudiantes de Bachillerato.
Responde SIEMPRE en {idioma}.

ESTRUCTURA: Si las imágenes contienen varios temas, organiza separando cada tema:
# Numero. Titulo del tema
[resumen]
---
# Numero. Titulo del tema
[resumen]

FORMATO:
- Título: # Numero. Titulo
- Apartados: ### 1. Nombre: explicación
- Conceptos clave: **concepto**

CALIDAD:
- COMPACTO: cada apartado en {lineas} líneas máximo
- Solo información esencial
- Lenguaje académico directo para memorizar en examen
- NO inventes información{sin_esquemas}{ctx}"""

    content = [{"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": encode_image(img)}} for img in imagenes]
    content.append({"type": "text", "text": "Analiza las imágenes y organiza el resumen por temas detectados automáticamente."})

    resp = client.messages.create(
        model=MODELO, max_tokens=3000, system=system,
        messages=[{"role": "user", "content": content}]
    )
    return resp.content[0].text

def extraer_titulo(resumen, contexto):
    if contexto.strip():
        return contexto.strip()[:60]
    for linea in resumen.split('\n'):
        linea = linea.strip()
        if linea.startswith('# '):
            return re.sub(r'[\\/*?:"<>|]', '', linea[2:].strip())[:60]
    return f"Resumen {datetime.datetime.now().strftime('%d/%m %H:%M')}"

def crear_docx(resumen):
    doc = Document()
    NEGRO = RGBColor(0, 0, 0)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Pt(56)
    doc.styles['Normal'].font.name = 'Liberation Serif'
    doc.styles['Normal'].font.size = Pt(12)

    def negritas(p, texto):
        for i, parte in enumerate(re.split(r'\*\*(.*?)\*\*', texto)):
            run = p.add_run(parte)
            run.font.name = 'Liberation Serif'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO
            if i % 2 == 1:
                run.bold = True

    for linea in resumen.split('\n'):
        linea = linea.strip()
        if not linea:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
        elif re.match(r'^# ', linea):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = Pt(17)
            run = p.add_run(linea[2:].strip())
            run.bold = True; run.underline = True
            run.font.name = 'Liberation Serif'; run.font.size = Pt(13); run.font.color.rgb = NEGRO
        elif re.match(r'^## ', linea):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(17)
            run = p.add_run(linea[3:].strip())
            run.bold = True; run.underline = True
            run.font.name = 'Liberation Serif'; run.font.size = Pt(12); run.font.color.rgb = NEGRO
        elif re.match(r'^### ', linea):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = Pt(17)
            texto_ap = linea[4:].strip()
            m = re.match(r'^([^:]+):(.*)', texto_ap)
            if m:
                run = p.add_run(m.group(1).strip() + ': ')
                run.bold = True; run.underline = True
                run.font.name = 'Liberation Serif'; run.font.size = Pt(12); run.font.color.rgb = NEGRO
                negritas(p, m.group(2).strip())
            else:
                run = p.add_run(texto_ap)
                run.bold = True; run.underline = True
                run.font.name = 'Liberation Serif'; run.font.size = Pt(12); run.font.color.rgb = NEGRO
        elif re.match(r'^[-] ', linea):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = Pt(17)
            negritas(p, linea[2:].strip())
        elif re.match(r'^\d+\. ', linea):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = Pt(17)
            negritas(p, re.sub(r'^\d+\. ', '', linea))
        elif linea == '---':
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = Pt(17)
            negritas(p, linea)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def guardar_en_historial(titulo, resumen):
    if "historial" not in st.session_state:
        st.session_state.historial = []
    st.session_state.historial.insert(0, {
        "titulo": titulo, "resumen": resumen,
        "fecha": datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    })
    if len(st.session_state.historial) > 20:
        st.session_state.historial = st.session_state.historial[:20]

def enviar_feedback(usuario, opinion, puntuacion):
    try:
        requests.post(FEEDBACK_URL, json={
            "fecha": datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),
            "usuario": usuario, "opinion": opinion, "puntuacion": str(puntuacion)
        }, timeout=5)
        return True
    except:
        return False

def mostrar_hero():
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "logo.png")
    if os.path.exists(logo_path):
        with open(logo_path, "rb") as f:
            logo_b64 = base64.b64encode(f.read()).decode()
        st.markdown(f'<div style="text-align:center;margin:0.5rem 0 0.3rem"><img src="data:image/png;base64,{logo_b64}" style="width:90px;height:90px;object-fit:contain;filter:drop-shadow(0 0 16px rgba(168,85,247,0.5));"></div>', unsafe_allow_html=True)
    st.markdown(f'<div class="hero-wrap"><div class="hero-badge">{t("badge")}</div><h1 class="hero-title">Shaad IA</h1><p class="hero-sub">{t("slogan")}</p></div>', unsafe_allow_html=True)

def check_password():
    if "autenticado" not in st.session_state:
        st.session_state.autenticado = False
    if not st.session_state.autenticado:
        mostrar_hero()
        st.markdown(f'<div class="card"><div class="card-label"><span class="dot"></span>{il("lock")} {t("login_badge")}</div><p style="color:var(--text-secondary);font-size:0.9rem;margin-bottom:1rem;">{t("login_desc")}</p></div>', unsafe_allow_html=True)
        password = st.text_input(t("contrasena"), type="password", placeholder=t("contrasena_ph"))
        if st.button(t("entrar"), type="primary"):
            if password == get_app_password():
                st.session_state.autenticado = True
                st.rerun()
            else:
                st.error(t("contrasena_error"))
        return False
    return True

if not check_password():
    st.stop()

mostrar_hero()

# Pestañas con iconos SVG inline usando markdown
tab_labels = [
    f"{icon('file-text', 'currentColor', 14)} {t('tab1')}",
    f"{icon('clock', 'currentColor', 14)} {t('tab2')}",
    f"{icon('settings', 'currentColor', 14)} {t('tab3')}",
    f"{icon('help-circle', 'currentColor', 14)} {t('tab4')}",
]
tab1, tab2, tab3, tab4 = st.tabs(tab_labels)

# ── TAB 1: RESUMIR ────────────────────────────────────────────────────────────
with tab1:
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(f'<p style="font-size:0.82rem;color:var(--text-secondary);margin-bottom:4px">{il("globe")} {t("idioma")}</p>', unsafe_allow_html=True)
        idioma = st.selectbox("idioma_sel", ["castellano", "galego", "català", "euskara", "english", "français", "português", "deutsch", "italiano"], label_visibility="collapsed")
    with col2:
        st.markdown(f'<p style="font-size:0.82rem;color:var(--text-secondary);margin-bottom:4px">{il("bar-chart")} {t("nivel")}</p>', unsafe_allow_html=True)
        nivel = st.selectbox("nivel_sel", ["Breve (2-3 líneas)", "Medio (4-6 líneas)", "Detallado (7-10 líneas)"], index=1, label_visibility="collapsed")
    with col3:
        st.markdown(f'<p style="font-size:0.82rem;color:var(--text-secondary);margin-bottom:4px">{il("table")} {t("tablas")}</p>', unsafe_allow_html=True)
        esquemas = st.toggle("tablas_tog", value=False, label_visibility="collapsed")

    st.markdown(f'<p style="color:var(--text-secondary);font-size:0.82rem;margin:0.8rem 0 0.4rem;">{il("message-square")} <strong style="color:var(--text-primary)">{t("instrucciones_titulo")}</strong></p>', unsafe_allow_html=True)

    if "contexto_val" not in st.session_state:
        st.session_state.contexto_val = ""

    chip_cols1 = st.columns(3)
    chip_cols2 = st.columns(3)
    for i, chip in enumerate(CHIPS):
        col = chip_cols1[i] if i < 3 else chip_cols2[i - 3]
        with col:
            if st.button(chip, key=f"chip_{i}"):
                st.session_state.contexto_val = chip
                st.rerun()

    contexto = st.text_input(
        "ctx",
        value=st.session_state.contexto_val,
        placeholder=t("instrucciones_ph"),
        label_visibility="collapsed"
    )
    if contexto != st.session_state.contexto_val:
        st.session_state.contexto_val = contexto

    # Subtexto encima del uploader (dentro de la caja lo gestiona Streamlit)
    st.markdown(f"""
    <div style="text-align:center;margin:1rem 0 -0.5rem">
        {icon("upload-cloud", "#a855f7", 36)}
        <p style="font-family:var(--font-display);font-size:1rem;font-weight:700;color:var(--text-primary);margin:0.3rem 0 0.1rem">{t("uploader_titulo")}</p>
        <p style="font-size:0.75rem;color:var(--text-muted);margin:0 0 0.5rem">{t("uploader_sub")}</p>
    </div>
    """, unsafe_allow_html=True)

    archivos = st.file_uploader(
        "uploader",
        type=["jpg", "jpeg", "png", "pdf"],
        accept_multiple_files=True,
        label_visibility="collapsed"
    )

    if archivos and len(archivos) > MAX_FOTOS:
        st.warning(t("aviso_max"))
        archivos = archivos[:MAX_FOTOS]

    imagenes_procesadas = []
    if archivos:
        for archivo in archivos:
            if archivo.type == "application/pdf":
                with st.spinner(f"PDF: {archivo.name}..."):
                    imagenes_procesadas.extend(pdf_a_imagenes(archivo))
            else:
                imagenes_procesadas.append(Image.open(archivo))
        if imagenes_procesadas:
            cols = st.columns(min(len(imagenes_procesadas), 4))
            for i, img in enumerate(imagenes_procesadas[:4]):
                with cols[i % 4]:
                    st.image(img, use_container_width=True)
            if len(imagenes_procesadas) > 4:
                st.caption(f"... +{len(imagenes_procesadas)-4}")

    st.markdown("<div style='margin-top:1.2rem'></div>", unsafe_allow_html=True)
    if st.button(t("btn_transformar"), type="primary"):
        if not imagenes_procesadas:
            st.warning(t("aviso_foto"))
        else:
            barra = st.progress(0, text=t("preparando"))
            barra.progress(20, text=t("analizando"))
            resultado = transcribir(imagenes_procesadas, idioma, nivel, contexto, esquemas)
            barra.progress(90, text=t("generando"))
            if resultado:
                barra.progress(100, text=t("listo"))
                titulo = extraer_titulo(resultado, contexto)
                guardar_en_historial(titulo, resultado)
                st.markdown(f'<div class="result-card"><span class="result-badge">✓ {titulo}</span></div>', unsafe_allow_html=True)
                st.markdown(resultado)
                docx_bytes = crear_docx(resultado)
                c1, c2 = st.columns(2)
                with c1:
                    st.download_button(t("descarga_txt"), data=resultado, file_name=f"{titulo}.txt", mime="text/plain", key="dl_txt")
                with c2:
                    st.download_button(t("descarga_docx"), data=docx_bytes, file_name=f"{titulo}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_docx")

# ── TAB 2: HISTORIAL ──────────────────────────────────────────────────────────
with tab2:
    historial = st.session_state.get("historial", [])
    if not historial:
        st.markdown(f"""
        <div class="empty-state">
            {icon("folder-open", "#a855f7", 48)}
            <p style="color:var(--text-secondary);font-size:0.95rem;margin:1rem 0 1.5rem;">{t("sin_resumenes")}</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button(t("ir_a_resumir"), type="primary"):
            st.rerun()
    else:
        for i, item in enumerate(historial):
            with st.expander(f"{item['titulo']} — {item['fecha']}"):
                st.markdown(item['resumen'])
                c1, c2 = st.columns(2)
                with c1:
                    st.download_button(t("descarga_txt"), data=item['resumen'], file_name=f"{item['titulo']}.txt", mime="text/plain", key=f"h_txt_{i}")
                with c2:
                    st.download_button(t("descarga_docx"), data=crear_docx(item['resumen']), file_name=f"{item['titulo']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"h_docx_{i}")

# ── TAB 3: AJUSTES ────────────────────────────────────────────────────────────
with tab3:
    st.markdown(f'<div class="card-label"><span class="dot"></span>{il("settings")} {t("ajustes_titulo")}</div>', unsafe_allow_html=True)

    st.markdown(f'<p style="font-size:0.85rem;color:var(--text-secondary);margin-bottom:4px">{t("tema_label")}</p>', unsafe_allow_html=True)
    tema = st.selectbox("tema", ["oscuro", "claro"],
        index=0 if get_tema() == "oscuro" else 1,
        format_func=lambda x: {"oscuro": "Oscuro", "claro": "Claro"}[x],
        label_visibility="collapsed")

    st.markdown(f'<p style="font-size:0.85rem;color:var(--text-secondary);margin:1rem 0 4px">{t("idioma_app_label")}</p>', unsafe_allow_html=True)
    idioma_app = st.selectbox("idioma_app", ["es", "en", "gl"],
        index=["es","en","gl"].index(get_idioma_app()),
        format_func=lambda x: {"es": "Castellano", "en": "English", "gl": "Galego"}[x],
        label_visibility="collapsed")

    st.markdown("<div style='margin-top:1rem'></div>", unsafe_allow_html=True)
    if st.button(t("guardar"), type="primary"):
        st.session_state.tema = tema
        st.session_state.idioma_app = idioma_app
        st.success(t("ajustes_guardados"))
        st.rerun()

# ── TAB 4: AYUDA ──────────────────────────────────────────────────────────────
with tab4:
    st.markdown(f'<div class="card-label"><span class="dot"></span>{il("help-circle")} Cómo usar Shaad IA</div>', unsafe_allow_html=True)

    pasos = [
        ("1", "Sube tus fotos o PDF", "Arrastra las imágenes o un PDF. Máximo 5 archivos."),
        ("2", "Elige idioma y nivel", "Selecciona el idioma del resumen y el nivel de detalle."),
        ("3", "Instrucciones rápidas", "Selecciona un chip o escribe tus propias instrucciones."),
        ("4", "Genera el resumen", "Pulsa el botón y espera unos segundos."),
        ("5", "Descarga", "Descarga en .txt o .docx con negritas y formato correctos."),
        ("·", "Consejo", "Si subes varias fotos de temas distintos, Shaad IA las organiza automáticamente por apartados."),
    ]
    for num, titulo_paso, desc in pasos:
        st.markdown(f'<div class="feedback-item"><strong style="color:var(--purple-bright)">{num}. {titulo_paso}</strong><br><span style="color:var(--text-secondary);font-size:0.9rem">{desc}</span></div>', unsafe_allow_html=True)

    st.markdown("<div style='margin-top:1.5rem'></div>", unsafe_allow_html=True)
    st.markdown(f'<div class="card-label"><span class="dot"></span>{il("user")} Tu nombre</div>', unsafe_allow_html=True)
    nombre_usuario = st.text_input("nombre", placeholder=t("nombre_ph"), label_visibility="collapsed", value=st.session_state.get("nombre_usuario", ""))
    if nombre_usuario:
        st.session_state.nombre_usuario = nombre_usuario

    st.markdown("<div style='margin-top:1rem'></div>", unsafe_allow_html=True)
    st.markdown(f'<div class="card-label"><span class="dot"></span>{il("message-square")} Deja tu opinión</div>', unsafe_allow_html=True)
    opinion = st.text_area("opinion", placeholder=t("opinion_ph"), label_visibility="collapsed")

    col_s1, col_s2 = st.columns([3, 1])
    with col_s1:
        puntuacion = st.select_slider("Puntuación (1 a 5)", options=[1, 2, 3, 4, 5], value=5)
    with col_s2:
        estrellas = "".join([icon("star", "#f9e2af", 14) for _ in range(puntuacion)])
        st.markdown(f'<div style="padding-top:1.8rem">{estrellas}</div>', unsafe_allow_html=True)

    if st.button(t("enviar_fb"), type="primary"):
        if opinion.strip():
            with st.spinner("Enviando…"):
                ok = enviar_feedback(st.session_state.get("nombre_usuario", "Anónimo"), opinion, puntuacion)
            st.success(t("fb_ok")) if ok else st.error(t("fb_error"))
        else:
            st.warning(t("fb_vacio"))

st.markdown(f'<div class="footer-custom">{t("footer")}</div>', unsafe_allow_html=True)